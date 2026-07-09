import streamlit as st
import docx
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor
import matplotlib.pyplot as plt
from io import BytesIO
import json
import os
import re

FAJL_BAZA = "baza_podatoci.json"

# --- МЕНАЏИРАЊЕ СО БАЗА ---
def vcitaj_podatoci():
    if os.path.exists(FAJL_BAZA):
        try:
            with open(FAJL_BAZA, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            pass
    return {
        "Математика 1": {
            "broj_poimi": 1,
            "poimi": [{
                "ime": "Матрица", 
                "def": "Правоаголна шема од броеви.", 
                "primeri": "А = [[1, 2], [3, 4]]", 
                "beleska": "Многу важна за линеарни трансформации."
            }],
            "smartart": ["Предавања", "Вежби", "Испит"],
            "glavi": {
                "Глава 1: Линеарна алгебра": {
                    "sodrzina": "**Важно:** Овде учиме за системи равенки.\n\n>> Прашање за селф-тест: Што е квадратна матрица?\n    Тоа е матрица каде бројот на редици е еднаков со бројот на колони (m = n).",
                    "status": "🟡 Се учи"
                }
            }
        }
    }

def zacuvaj_podatoci():
    with open(FAJL_BAZA, "w", encoding="utf-8") as f:
        json.dump(st.session_state.ispiti, f, ensure_ascii=False, indent=4)

# --- ПОМОШНИ ФУНКЦИИ ЗА WORD ---
def make_heading_collapsed_by_default(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    collapsed = OxmlElement('w:collapsed')
    collapsed.set(qn('w:val'), 'true')
    pPr.append(collapsed)

def add_hyperlink_to_bookmark(paragraph, text, bookmark_name, color="0000FF"):
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="{bookmark_name}"/>')
    new_run = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:rColor w:val="{color}"/><w:u w:val="single"/></w:rPr><w:t>{text}</w:t></w:r>')
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_bookmark(paragraph, bookmark_name):
    p = paragraph._p
    id_num = "1"
    start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{id_num}" w:name="{bookmark_name}"/>')
    end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{id_num}"/>')
    p.insert(0, start)
    p.append(end)

def generate_smartart_image(steps):
    fig, ax = plt.subplots(figsize=(6, 2))
    ax.axis('off')
    for i, step in enumerate(steps):
        ax.text(i*2.5 + 1, 0.5, step, ha='center', va='center', color='white', weight='bold',
                bbox=dict(boxstyle='round,pad=0.6', facecolor='#2383E2', edgecolor='none'))
        if i < len(steps) - 1:
            ax.text(i*2.5 + 2.25, 0.5, '➔', ha='center', va='center', fontsize=20, color='#7c7b77')
    ax.set_xlim(0, len(steps)*2.5)
    ax.set_ylim(0, 1)
    img_buf = BytesIO()
    plt.savefig(img_buf, format='png', bbox_inches='tight', dpi=150)
    img_buf.seek(0)
    plt.close()
    return img_buf

def add_callout_box(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="2383E2"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    cell._tc.get_or_add_tcPr().append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.1)
    return p

def procesiraj_tekst_so_stilovi(paragraph, tekst, recnik_poimi, избран_испит):
    tokens = re.split(r'(\*\*.*?\*\*|==.*?==|\^.*?\^)', tekst)
    for token in tokens:
        if not token:
            continue
        is_bold = token.startswith('**') and token.endswith('**')
        is_high = token.startswith('==') and token.endswith('==')
        is_super = token.startswith('^') and token.endswith('^')
        
        чист_дел = token
        if is_bold: чист_дел = token[2:-2]
        elif is_high: чист_дел = token[2:-2]
        elif is_super: чист_дел = token[1:-1]
        
        зборови = чист_дел.split(" ")
        for и, збор in enumerate(зборови):
            последен = (и == len(зборови) - 1)
            додаток = "" if последен else " "
            клуч_збор = збор.strip(",.!?\"'")
            
            if клуч_збор in recnik_poimi:
                add_hyperlink_to_bookmark(paragraph, збор + додаток, f"ref_{избран_испит}_{клуч_збор}")
            else:
                run = paragraph.add_run(збор + додаток)
                if is_bold: run.bold = True
                if is_super: run.font.superscript = True
                if is_high:
                    highlight_xml = parse_xml(f'<w:highlight {nsdecls("w")} w:val="yellow"/>')
                    run._r.get_or_add_rPr().append(highlight_xml)

# --- НОВА ФУНКЦИЈА: ПРИКАЗ НА ТЕКСТ ВО ЖИВО НА ЕКРАНОТ СО ПОДДРШКА ЗА TOGGLE ---
def prikazi_sodrzina_na_ekran(tekst_sodrzina):
    redovi = tekst_sodrzina.split('\n')
    aktivno_toggle_meni = None
    toggle_vnatresen_tekst = []

    for ред in redovi:
        # 1. Спречуваме грешки со празни редови
        if not ред.strip():
            if aktivno_toggle_meni:
                toggle_vnatresen_tekst.append("")
            else:
                st.write("")
            continue

        # 2. Ако најдеме новToggle наслов (започнува со >>)
        if ред.strip().startswith('>>'):
            # Прво го затвораме претходниот отворен toggle ако имало таков
            if aktivno_toggle_meni:
                with aktivno_toggle_meni:
                    st.markdown("\n".join(toggle_vnatresen_tekst))
                toggle_vnatresen_tekst = []
            
            naslov = ред.strip()[2:].strip()
            # Креираме визуелно Toggle мени на самиот екран во Streamlit
            aktivno_toggle_meni = st.expander(f"📐 {naslov}", expanded=False)
            continue

        # 3. Ако редот има 4 празни места (ова е содржина што треба да оди внатре во Toggle)
        if ред.startswith('    ') or ред.startswith('\t'):
            if aktivno_toggle_meni:
                # Ја чистиме замена за маркер (==текст== -> во Streamlit се заменува со HTML/Markdown приказ ако е потребно, но за едноставност користиме чист маркер преку Streamlit markdown)
                исчистен_текст = ред.strip()
                # Краток фикс за маркерот == да работи во Streamlit со чисто HTML бидејќи стандардниот markdown не поддржува ==
                исчистен_текст = re.sub(r'==(.*?)==', r'<mark style="background-color: yellow; color: black;">\1</mark>', исчистен_текст)
                toggle_vnatresen_tekst.append(исчистен_текст)
            else:
                st.markdown(ред.strip(), unsafe_allow_html=True)
            continue

        # 4. Ако наидеме на обичен текст, го затвораме моменталниот Toggle
        if aktivno_toggle_meni:
            with aktivno_toggle_meni:
                st.markdown("\n".join(toggle_vnatresen_tekst), unsafe_allow_html=True)
            aktivno_toggle_meni = None
            toggle_vnatresen_tekst = []

        # Обични стилови на екранот
        исчистен_ред = ред.strip()
        исчистен_ред = re.sub(r'==(.*?)==', r'<mark style="background-color: yellow; color: black;">\1</mark>', исчистен_ред)
        
        if исчистен_ред.startswith('>'):
            st.info(исчистен_ред[1:].strip())
        else:
            st.markdown(исчистен_ред, unsafe_allow_html=True)

    # На крајот, ако останало отворено последното toggle мени, го исцртуваме
    if aktivno_toggle_meni:
        with aktivno_toggle_meni:
            st.markdown("\n".join(toggle_vnatresen_tekst), unsafe_allow_html=True)

# --- ИНИЦИЈАЛИЗАЦИЈА ---
if "ispiti" not in st.session_state:
    st.session_state.ispiti = vcitaj_podatoci()

st.set_page_config(page_title="Notion Workspace за Испити", layout="centered")

# --- СТРАНИЧНО МЕНИ (SIDEBAR) ---
st.sidebar.title("📚 Notion Училница")

нов_испит = st.sidebar.text_input("Име на нов испит:")
if st.sidebar.button("➕ Креирај испит"):
    if нов_испит.strip() and нов_испит not in st.session_state.ispiti:
        st.session_state.ispiti[нов_испит.strip()] = {
            "broj_poimi": 1,
            "poimi": [{"ime": "", "def": "", "primeri": "", "beleska": ""}],
            "smartart": ["Чекор 1", "Чекор 2", "Чекор 3"],
            "glavi": {"Глава 1": {"sodrzina": "", "status": "🔴 Недопрено"}}
        }
        zacuvaj_podatoci()
        st.rerun()

st.sidebar.markdown("---")

if st.session_state.ispiti:
    листа_испити = list(st.session_state.ispiti.keys())
    избран_испит = st.sidebar.selectbox("Избери предмет:", листа_испити)
    podatoci = st.session_state.ispiti[избран_испит]
    
    нова_глава_име = st.sidebar.text_input("Име на нова Глава/Страница:")
    if st.sidebar.button("📂 Додај нова глава"):
        if нова_глава_име.strip() and нова_глава_име.strip() not in podatoci["glavi"]:
            podatoci["glavi"][нова_глава_име.strip()] = {"sodrzina": "", "status": "🔴 Недопрено"}
            zacuvaj_podatoci()
            st.rerun()
            
    листа_глави = list(podatoci["glavi"].keys())
    избрана_глава = st.sidebar.radio("Избери активна Глава:", листа_глави)
    
    ново_име_глава = st.sidebar.text_input("Преименувај ја главата:", value=избрана_глава)
    if st.sidebar.button("💾 Зачувај име на глава"):
        if ново_име_глава.strip() and ново_име_глава.strip() != избрана_глава:
            podatoci["glavi"][ново_име_глава.strip()] = podatoci["glavi"].pop(избрана_глава)
            zacuvaj_podatoci()
            st.rerun()
            
    ново_име = st.sidebar.text_input("Преименувај го предметот:", value=избран_испит)
    if st.sidebar.button("💾 Зачувај име на предмет"):
        if ново_име.strip() and ново_име.strip() != избран_испит:
            st.session_state.ispiti[ново_име.strip()] = st.session_state.ispiti.pop(избран_испит)
            zacuvaj_podatoci()
            st.rerun()
            
    if st.sidebar.button("🗑️ Избриши го целиот предмет", type="primary", key="del_predmet"):
        st.session_state.ispiti.pop(избран_испит)
        zacuvaj_podatoci()
        st.rerun()

st.sidebar.markdown("---")
st.sidebar.subheader("🔄 Синхронизација (Локално/Cloud)")
baza_json_str = json.dumps(st.session_state.ispiti, ensure_ascii=False, indent=4)
st.sidebar.download_button(label="📥 Спушти ја базата (.json)", data=baza_json_str, file_name="baza_podatoci.json", mime="application/json")

# --- ГЛАВЕН ЕКРАН ---
if st.session_state.ispiti:
    st.title(f"📖 {избран_испит}")
    
    col_h, col_s = st.columns([2, 1])
    with col_h:
        st.subheader(f"📄 Глава: {избрана_глава}")
    with col_s:
        тековен_status = podatoci["glavi"][избрана_глава].get("status", "🔴 Недопрено")
        избран_status = st.selectbox("Статус:", ["🔴 Недопрено", "🟡 Се учи", "🟢 Научено"], index=["🔴 Недопрено", "🟡 Се учи", "🟢 Научено"].index(тековен_status))
        if избран_status != тековен_status:
            podatoci["glavi"][избрана_глава]["status"] = избран_status
            zacuvaj_podatoci()

    # Поле за пишување содржина
    тековна_содржина = podatoci["glavi"][избрана_глава]["sodrzina"]
    nova_sodrzina = st.text_area("Внесете ја содржината за оваа глава:", value=тековна_содржина, height=200, key=f"area_{избран_испит}_{избрана_глава}")
    if nova_sodrzina != тековна_содржина:
        podatoci["glavi"][избрана_глава]["sodrzina"] = nova_sodrzina
        zacuvaj_podatoci()

    # ✨ ОВА Е КЛУЧНИОТ ДЕЛ: ПРЕГЛЕД ВО ЖИВО НА ЕКРАНОТ СО ИДЕНТИЧЕН TOGGLE ЕФЕКТ
    st.markdown("### 👀 Преглед на лекцијата во живо (Notion Стил):")
    prikazi_sodrzina_na_ekran(nova_sodrzina)

    st.markdown("---")

    # --- ГЛОБАЛЕН РЕЧНИК ---
    st.subheader("📘 Глобален Речник на поими")
    recnik_poimi = {}
    novi_poimi = []
    
    for p_idx in range(podatoci["broj_poimi"]):
        default_ime = podatoci["poimi"][p_idx].get("ime", "")
        default_def = podatoci["poimi"][p_idx].get("def", "")
        default_primeri = podatoci["poimi"][p_idx].get("primeri", "")
        default_beleska = podatoci["poimi"][p_idx].get("beleska", "")
        
        c1, c2, c3, c4 = st.columns([1.2, 2, 1.5, 1.5])
        with c1: p_ime = st.text_input("Поим:", key=f"{избран_испит}_p_ime_{p_idx}", value=default_ime)
        with c2: p_def = st.text_input("Дефиниција:", key=f"{избран_испит}_p_def_{p_idx}", value=default_def)
        with c3: p_primeri = st.text_input("Примери:", key=f"{избран_испит}_p_prim_{p_idx}", value=default_primeri)
        with c4: p_beleska = st.text_input("Белешка:", key=f"{избран_испит}_p_bel_{p_idx}", value=default_beleska)
            
        novi_poimi.append({"ime": p_ime, "def": p_def, "primeri": p_primeri, "beleska": p_beleska})
        if p_ime.strip(): recnik_poimi[p_ime.strip()] = p_def

    if novi_poimi != podatoci["poimi"]:
        podatoci["poimi"] = novi_poimi
        zacuvaj_podatoci()

    if st.button("➕ Додај нов поим во речникот"):
        podatoci["broj_poimi"] += 1
        zacuvaj_podatoci()
        st.rerun()

    st.markdown("---")
    # SMARTART ПЛАН ЗА УЧЕЊЕ
    st.subheader("📊 Процес на учење")
    cekor1 = st.text_input("Чекор 1:", key=f"{избран_испит}_c1", value=podatoci["smartart"][0])
    cekor2 = st.text_input("Чекор 2:", key=f"{избран_испит}_c2", value=podatoci["smartart"][1])
    cekor3 = st.text_input("Чекор 3:", key=f"{избран_испит}_c3", value=podatoci["smartart"][2])
    if [cekor1, cekor2, cekor3] != podatoci["smartart"]:
        podatoci["smartart"] = [cekor1, cekor2, cekor3]
        zacuvaj_podatoci()

    st.markdown("---")

    # ГЕНЕРИРАЊЕ WORD СКУПШТИНА
    if st.button(f"🎬 Креирај комплетна скрипта за {избран_испит}"):
        doc = Document()
        doc.add_heading(f'Скрипта за испит: {избран_испит}', level=0)
        
        doc.add_heading('📊 План и чекори за учење', level=1)
        img_data = generate_smartart_image(podatoci["smartart"])
        doc.add_picture(img_data, width=Inches(5.5))
        doc.add_paragraph()
        
        for наслов_глава, инфо_глава in podatoci["glavi"].items():
            heading = doc.add_heading(наслов_глава, level=1)
            make_heading_collapsed_by_default(heading)
            
            p_status = doc.add_paragraph()
            p_status.add_run(f"Статус: {инфо_глава.get('status', '🔴 Недопрено')}\n").italic = True
            
            редови = инфо_глава["sodrzina"].split('\n')
            for ред in редови:
                if not ред.strip(): continue
                if ред.strip().startswith('>>'):
                    p_toggle = doc.add_heading(ред.strip()[2:].strip(), level=3)
                    make_heading_collapsed_by_default(p_toggle)
                elif ред.startswith('    ') or ред.startswith('\t'):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.4)
                    procesiraj_tekst_so_stilovi(p, ред.strip(), recnik_poimi, избран_испит)
                elif ред.strip().startswith('>'):
                    p = add_callout_box(doc, ред.strip()[1:].strip())
                    procesiraj_tekst_so_stilovi(p, ред.strip()[1:].strip(), recnik_poimi, избран_испит)
                elif ред.strip().startswith('* ') or ред.strip().startswith('- '):
                    p = doc.add_paragraph(style='List Bullet')
                    procesiraj_tekst_so_stilovi(p, ред.strip()[2:], recnik_poimi, избран_испит)
                elif re.match(r'^\d+\.\s', ред.strip()):
                    p = doc.add_paragraph(style='List Number')
                    procesiraj_tekst_so_stilovi(p, re.sub(r'^\d+\.\s', '', ...), recnik_poimi, избран_испит)
                else:
                    p = doc.add_paragraph()
                    procesiraj_tekst_so_stilovi(p, ред, recnik_poimi, избран_испит)
                    
        doc.add_heading('📘 Заеднички речник на поими', level=1)
        for p_struktura in podatoci["poimi"]:
            poim = p_struktura["ime"].strip()
            if not poim: continue
            p_poim = doc.add_paragraph()
            add_bookmark(p_poim, f"ref_{избран_испит}_{poim}")
            p_poim.add_run(f"• {poim}: ").bold = True
            p_poim.add_run(p_struktura["def"])
            
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        st.success("🎉 Скриптата е генерирана!")
        st.download_button(label="📥 Преземи .docx скрипта", data=bio, file_name=f"Скрипта_{избран_испит}.docx")